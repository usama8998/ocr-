from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pytesseract
from PIL import Image
from docx import Document
import os

app = Flask(__name__)
CORS(app)  # Allow cross-origin requests

# Tesseract ka path set karo
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"  # Render par default path

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'images' not in request.files:
        return jsonify({"error": "No images uploaded"}), 400

    files = request.files.getlist("images")
    if not files:
        return jsonify({"error": "No files found"}), 400

    files.sort(key=lambda x: x.filename)  # Sort files by name to maintain page order

    doc = Document()

    for idx, file in enumerate(files):
        try:
            image = Image.open(file)
        except Exception as e:
            return jsonify({"error": f"Error processing file {file.filename}: {str(e)}"}), 500

        # OCR se text nikalo
        extracted_text = pytesseract.image_to_string(image, lang='eng+urd')

        # Word document mein text daalo
        doc.add_heading(f"Page {idx+1}", level=2)
        doc.add_paragraph(extracted_text)

        if idx < len(files) - 1:
            doc.add_page_break()

    output_path = "/tmp/MergedPaper.docx"
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0", port=10000)
